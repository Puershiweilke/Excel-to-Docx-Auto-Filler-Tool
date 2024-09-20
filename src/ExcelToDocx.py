import pandas as pd
import docx
from docx import Document
from docx.shared import Inches
import tkinter as tk
from tkinter import filedialog, messagebox
import os
import re
import requests

class Fake_Doc_obj:
    def __init__(self, text):
        self.text = text

# obj 是 doc对象
# df 是 dataframe 是pandas的一个数据结构
def replace_of_template(obj, df,img_size=1,type="paragraph"):
    # 定义图片链接的正则表达式
    image_link_pattern = re.compile(r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\$\$,]|(?:%[0-9a-fA-F][0-9a-fA-F]))+.(?:jpg|jpeg|png|gif)$')
    suffix_r = re.compile(r'(jpg|jpeg|png|gif)$')
    for key in df.keys():
        value = str(df.loc[key])
        text = obj.text.replace("{{" + key + "}}", value)
        if text != obj.text:
            # 使用正则表达式判断它是不是图片链接
            if image_link_pattern.search(value):
                ##获取图片的后缀
                suffix = suffix_r.search(value).group(0)
                ##删除占位符
                obj.text = obj.text.replace("{{" + key + "}}","")
                # 如果是图片链接，可以在这里进行进一步的处理
                try:
                    ##尝试下载图片，并插入到doc相应的位置
                    response = requests.get(value)
                    response.raise_for_status()
                    img_data = response.content

                    # 保存图片到本地
                    with open(f"temp.{suffix}", "wb") as img_file:
                        img_file.write(img_data)

                    img_path = f"./temp.{suffix}"
                    if type == "cell":
                        # 插入图片
                        obj.add_paragraph().add_run().add_picture(img_path,width=Inches(img_size))

                    elif type == "paragraph":
                        # 插入图片
                        obj.add_run().add_picture(img_path,width=Inches(img_size))
                    else:
                        raise  AttributeError("错误：在填入图片时遇到了表格、段落以外的容器。")
                    print(f"Image downloaded and added to document: {value}")
                    os.remove(img_path)
                except requests.RequestException as e:
                    print(f"Failed to download image: {value}. Error: {e}")
            else:
                obj.text = text


    return obj.text

def generate_documents(excel_path, template_path, output_folder, filename_template,img_size=1):
    # 读取 Excel 文件
    df = pd.read_excel(excel_path)

    # 遍历每一行数据
    for index, row in df.iterrows():
        # 读取 DOCX 模板
        doc = Document(template_path)
        #循环doc的所有数据，然后替换
        #首先处理段落
        for para in doc.paragraphs:
            replace_of_template(para,df.iloc[index],img_size,"paragraph")

        #然后处理表格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    replace_of_template(cell,df.iloc[index],img_size,"cell")

        #处理文件名
        filename = replace_of_template(Fake_Doc_obj(filename_template),df.iloc[index])
        output_path = os.path.join(output_folder,filename+".docx")
        # 保存生成的文档
        print(output_path)
        doc.save(output_path)
    
    messagebox.showinfo("完成", "文档生成完成！")

def select_file(entry):
    filename = filedialog.askopenfilename()
    entry.delete(0, tk.END)
    entry.insert(0, filename)

def select_folder(entry):
    foldername = filedialog.askdirectory()
    entry.delete(0, tk.END)
    entry.insert(0, foldername)

def run_generation():
    excel_path = excel_entry.get()
    template_path = template_entry.get()
    output_folder = output_entry.get()
    filename_template = filename_template_entry.get()
    img_size = float(img_size_entry.get())

    generate_documents(excel_path, template_path, output_folder, filename_template,img_size)

# 创建主窗口
root = tk.Tk()
root.title("从excel到docx的自动填写小工具 An Automatic Filler from Excel to DOCX")

# 输入框和标签
tk.Label(root, text="Excel 文件路径(excel file path)").grid(row=1)
tk.Label(root, text="DOCX 模板路径(docx template path)").grid(row=2)
tk.Label(root, text="输出文件夹路径 output path").grid(row=3)
tk.Label(root, text="文件名模板 (使用模板语法 例：{{key}})(name template)").grid(row=4)
tk.Label(root, text="图片尺寸（英寸）(不需要的话设置为 1 即可)(image size(inche) )").grid(row=5)

excel_entry = tk.Entry(root, width=50)
template_entry = tk.Entry(root, width=50)
output_entry = tk.Entry(root, width=50)
filename_template_entry = tk.Entry(root, width=50)
img_size_entry = tk.Entry(root, width=50)

excel_entry.grid(row=1, column=1)
template_entry.grid(row=2, column=1)
output_entry.grid(row=3, column=1)
filename_template_entry.grid(row=4, column=1)
img_size_entry.grid(row=5, column=1)

tk.Button(root, text="选择文件(choice file)", command=lambda: select_file(excel_entry)).grid(row=1, column=2)
tk.Button(root, text="选择模板(choice template)", command=lambda: select_file(template_entry)).grid(row=2, column=2)
tk.Button(root, text="选择输出文件夹(choice output folder )", command=lambda: select_folder(output_entry)).grid(row=3, column=2)

# 提交按钮
tk.Button(root, text="生成docx文档(generate docx document)", command=run_generation).grid(row=6, columnspan=3)

# 启动 GUI
root.mainloop()
